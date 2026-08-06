# -*- coding: utf-8 -*-
"""
Full research paper for Paper 1 following Neha Dhiman template structure.
ONLY genuine multi-model results (Optimized/metrics_fixed.py). No COM_A / mealpy fabrications.
"""
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\USER\Downloads\PostDoc\CODE_28-04-2025_Paper1")
OPT = ROOT / "Optimized"
RES = OPT / "results"
FIGS = OPT / "figures"
OUT = OPT / "Paper1_SMA_CLMPNet_Genuine_Research_Paper.docx"
OUT2 = ROOT / "Paper1_SMA_CLMPNet_Genuine_Research_Paper.docx"


def set_run(run, size=11, bold=False, italic=False, color=None, super=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if super:
        run.font.superscript = True


def P(doc, text="", size=11, bold=False, italic=False, align="left",
      before=0, after=6, first=0, color=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15
    if first:
        pf.first_line_indent = Inches(first)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def H(doc, text, level=1):
    sizes = {1: 13, 2: 12, 3: 11}
    p = P(doc, "", size=sizes.get(level, 11), before=12 if level == 1 else 8, after=4)
    r = p.add_run(text)
    set_run(r, size=sizes.get(level, 11), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p


def body(doc, text):
    return P(doc, text, size=11, align="justify", after=8, first=0.3)


def cap(doc, text):
    return P(doc, text, size=9, italic=True, align="center", before=2, after=10)


def shade(cell, hx):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hx)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run(r, size=8, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(cell, "1F4E79")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run(r, size=8)
            if i % 2:
                shade(cell, "E8F0F8")
    P(doc, "", after=4)
    return t


def add_img(doc, path, width=5.5, caption=None):
    path = Path(path)
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cap(doc, caption)


def read_csv(path):
    out = {}
    path = Path(path)
    if not path.exists():
        return out
    with open(path, newline="", encoding="utf-8") as fh:
        for row in csv.DictReader(fh):
            m = row["model"]
            tp = int(float(row["train_pct"]))
            out.setdefault(m, {})[tp] = row
    return out


def pct(v):
    try:
        return f"{float(v) * 100:.2f}"
    except Exception:
        return "—"


def collapse_mark(r):
    try:
        sen, spe = float(r.get("SEN", 0)), float(r.get("SPE", 0))
        if (sen < 0.05 and spe > 0.95) or (spe < 0.05 and sen > 0.95):
            return "†"
    except Exception:
        pass
    return ""


def metrics_table(doc, data, models, title):
    H(doc, title, 3)
    P(
        doc,
        "GENUINE results only (sklearn metrics_fixed). Values in %. "
        "† = class collapse (prefer BalAcc/F1 over Acc).",
        size=9, italic=True, after=4,
    )
    rows = []
    for m in models:
        r80 = data.get(m, {}).get(80, {})
        r90 = data.get(m, {}).get(90, {})
        m80 = m + collapse_mark(r80)
        m90 = collapse_mark(r90)
        rows.append([
            m80,
            pct(r80.get("ACC")), pct(r80.get("SEN")), pct(r80.get("SPE")),
            pct(r80.get("PRE")), pct(r80.get("F1")), pct(r80.get("BAL_ACC")),
            pct(r90.get("ACC")) + m90, pct(r90.get("SEN")), pct(r90.get("SPE")),
            pct(r90.get("PRE")), pct(r90.get("F1")), pct(r90.get("BAL_ACC")),
        ])
    table(
        doc,
        ["Model", "A80", "Se80", "Sp80", "P80", "F80", "B80",
         "A90", "Se90", "Sp90", "P90", "F90", "B90"],
        rows,
    )


def author_block(doc):
    p = P(doc, "", align="center", after=2)
    parts = [
        ("Dr. Abhishek Thakur", {"size": 12, "bold": True}),
        ("1,2,*", {"size": 9, "super": True}),
        (",  ", {"size": 12}),
        ("Prof. Vishal Jain", {"size": 12, "bold": True}),
        ("3", {"size": 9, "super": True}),
        (", and  ", {"size": 12}),
        ("Prof. Chin-Shiuh Shieh", {"size": 12, "bold": True}),
        ("4", {"size": 9, "super": True}),
    ]
    for text, kw in parts:
        r = p.add_run(text)
        set_run(r, **kw)
    P(doc, "(謝欽旭)", size=9, italic=True, align="center", after=8)

    affs = [
        ("1", " Primary: School of Computer Science & Engineering, Chitkara University, Himachal Pradesh, India (Associate Professor)."),
        ("2", " Secondary: Postdoctoral Researcher, RIITC / Department of Electronic Engineering, NKUST, Kaohsiung 80778, Taiwan (R.O.C.); Ref. RIITC-Postdoc-2027-B03."),
        ("3", " Co-Supervisor: Vivekananda Institute of Professional Studies – Technical Campus (VIPS-TC), India."),
        ("4", " Supervisor: Department of Electronic Engineering / RIITC, NKUST, No. 415, Jiangong Rd., Kaohsiung 80778, Taiwan (R.O.C.)."),
    ]
    for sup, text in affs:
        p = P(doc, "", align="center", after=2)
        r = p.add_run(sup)
        set_run(r, size=8, super=True)
        r2 = p.add_run(text)
        set_run(r2, size=9, italic=True)

    P(doc, "E-mail: abhithakur25@gmail.com; abhishek@chitkarauniversity.edu.in  |  "
           "drvishaljain83@gmail.com; vishal.jain@vips.edu  |  csshieh@nkust.edu.tw",
      size=8, align="center", after=4)
    P(doc, "* Corresponding author: Dr. Abhishek Thakur (abhithakur25@gmail.com).",
      size=9, italic=True, align="center", after=10)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # Title (template-aligned)
    P(
        doc,
        "SMA-CLMPNet: Spatial Multiscale Attention enabled Convolutional Distributed "
        "Memory Network for Intra-frame Video Forgery Detection — A Machine-Verified "
        "Multi-Model Study with Genuine Metrics",
        size=15, bold=True, align="center", after=12,
    )
    author_block(doc)

    # Abstract
    H(doc, "Abstract", 1)
    P(
        doc,
        "Video forgery detection is essential for protecting digital media authenticity against "
        "deepfakes, splicing, and related identity-preserving manipulations. This paper presents "
        "SMA-CLMPNet (Spatial Multiscale Attention coupled Convolutional Distributed LSTM-based "
        "Modified Pooling Network) for passive intra-frame video forgery detection, combining "
        "gradient-based frame selection, Viola–Jones ROI extraction, multi-stream features "
        "(Grad-CAM, ResNet-101 statistics, VGG16-LDZP, optical flow), and a 3D-CNN with distributed "
        "LSTM plus SCAM/MUSE multiscale attention. Following a Paper-2-style multi-model protocol, "
        "we re-evaluate DCNN, EfficientNetV2B0, MobileNetV2, STIDNet, and the proposed P1 model on "
        "FaceForensics++ FaceSwap feature tensors of shape (50, 10, 128, 128, 12) using only "
        "honest sklearn metrics (Optimized/metrics_fixed.py). Fabricated mealpy-based Analysis "
        "scores are excluded. On this 50-sample cache (29/21 labels; majority baseline 58%), "
        "genuine machine re-runs at 20 epochs yield best non-collapse accuracy of 80% "
        "(DCNN @90% train, F1=0.67, BalAcc=0.75) and best balanced accuracy of 81.3% on prior "
        "balanced runs; the fresh 2026-08-06 balanced re-run peaks at BalAcc 68.8% "
        "(EfficientNetV2B0 @80%). Results demonstrate transparent multi-model comparison under "
        "integrity-safe scoring and align with RIITC–NKUST postdoctoral Objective 1 on optimised "
        "unimodal visual forensic extractors.",
        size=10, align="justify", after=8,
    )
    P(
        doc,
        "Keywords: Digital forensics; intra-frame video forgery; SMA-CLMPNet; SCAM; MUSE; "
        "Bi-LSTM / distributed LSTM; EfficientNetV2B0; MobileNetV2; FaceForensics++; genuine metrics.",
        size=10, italic=True, after=12,
    )

    # 1 Intro
    H(doc, "1. Introduction", 1)
    body(
        doc,
        "Detecting video forgery is essential for safeguarding digitized video footage and "
        "guaranteeing the validity of visual evidence across journalism, law enforcement, "
        "criminal investigation, medical imaging, insurance, and online platforms. Advanced "
        "editing tools and generative models enable deepfakes, splicing, and face-swap "
        "manipulations that are difficult to detect by human inspection alone.",
    )
    body(
        doc,
        "Classical approaches include block-based and keypoint-based copy-move detectors, "
        "statistical noise methods, and early deep CNNs. Block methods raise computational cost "
        "on overlapping tiles; keypoint methods underperform on smooth regions. Deep models "
        "improve accuracy but often lack transparent multi-architecture comparison under the "
        "same feature tensors and honest evaluation code paths.",
    )
    body(
        doc,
        "This work focuses on intra-frame forgery (spatial manipulation inside frames), "
        "especially FaceSwap-style facial replacement in FaceForensics++. The proposed "
        "SMA-CLMPNet integrates Spatial Multiscale Attention (SCAM + MUSE) with a convolutional "
        "distributed LSTM and modified pooling. The experimental section reports only "
        "machine-verified multi-model results scored with sklearn (not the vendored mealpy "
        "metric implementation, which has been shown to corrupt predictions).",
    )
    body(
        doc,
        "Contributions: (i) complete SMA-CLMPNet pipeline description aligned with the Paper 1 "
        "template; (ii) genuine multi-model cohort (DCNN, EfficientNetV2B0, MobileNetV2, STIDNet, "
        "P1-Proposed); (iii) class-weight / oversample protocol; (iv) exclusion of fabricated "
        "Analysis/COM_A artefacts from the repository and manuscript tables.",
    )

    # 2 Lit
    H(doc, "2. Literature Review", 1)
    body(
        doc,
        "Prior video forgery work spans STIDNet-style identity/spatio-temporal networks, DeepCNN "
        "object-based detectors, sensor pattern noise for splicing, EfficientNet multi-feature "
        "fusion with attention, 3D-CNNs for temporal localization, GLCM frame-duplication "
        "detectors, and CNN models with spatial–temporal attention for face video integrity. "
        "Each line has strengths but also challenges: sensitivity to post-processing, weak "
        "generalization to high-quality deepfakes, computational cost on high-resolution frames, "
        "and limited detection of facial inter-frame structure changes.",
    )
    H(doc, "2.1 Challenges", 2)
    for ch in [
        "Facial forgeries after heavy post-processing reduce DCNN reliability.",
        "Noise-pattern methods miss small/local alterations when SPN is weak.",
        "Deepfake diversity reduces generalization of single-backbone detectors.",
        "High-resolution 3D pipelines raise cost and slow feature refinement.",
        "Texture matrices (e.g., GLCM) may miss facial-structure manipulations.",
        "Evaluation integrity: some code paths report inflated metrics unrelated to model outputs.",
    ]:
        P(doc, "• " + ch, size=11, after=3)

    H(doc, "2.2 Problem Statement", 2)
    body(
        doc,
        "Let a labelled video collection provide authentic and forged clips. Frames are selected "
        "and pre-processed to ROIs; multi-stream features are fused and classified. The learning "
        "objective uses categorical cross-entropy between true labels and predicted class "
        "probabilities. The operational goal is accurate passive detection of intra-frame forgery "
        "with transparent, reproducible metrics under limited local feature caches.",
    )

    # 3 System model
    H(doc, "3. System Model", 1)
    body(
        doc,
        "The intra-frame video forgery system determines whether a clip is authentic or forged. "
        "Videos are collected, frames selected by gradient energy, faces extracted (Viola–Jones / "
        "Haar cascade), multi-features computed, and a deep model with multiscale attention "
        "produces a binary decision. Figure assets in the original template illustrate the "
        "end-to-end flow; this machine package stores pipeline screenshots under driver_out/ "
        "when the GUI driver is executed.",
    )

    # 4 Proposed
    H(doc, "4. Proposed SMA-CLMPNet for Intra-frame Video Forgery Detection", 1)
    body(
        doc,
        "SMA-CLMPNet improves feature discrimination using Spatial Multiscale Attention (SMA) that "
        "couples Spatial-Channel Joint Attention (SCAM) with Multiscale Excited (MUSE) attention, "
        "and a convolutional distributed LSTM with modified pooling (average top-k and log-sum-exp "
        "pooling blended by a balance weight).",
    )

    H(doc, "4.1 Collection and Preprocessing", 2)
    body(
        doc,
        "Videos are drawn from FaceForensics++ FaceSwap (and authentic youtube) sequences. "
        "Gradient-based frame selection retains high-edge-content frames; Viola–Jones ROI "
        "extraction focuses on face regions for subsequent feature extractors.",
    )

    H(doc, "4.2 Multi-feature Extraction", 2)
    body(
        doc,
        "Four complementary streams are extracted: (i) Grad-CAM importance heatmaps for "
        "class-discriminative regions; (ii) ResNet-101 residual embeddings summarized by "
        "statistical moments (mean, std, variance, skewness, kurtosis); (iii) VGG16 hierarchical "
        "maps combined with Local Directional ZigZag Pattern (LDZP) texture codes; (iv) "
        "Lucas–Kanade optical-flow motion fields. Streams are stacked into a multi-channel "
        "volume for 3D convolution. In the local Optimized evaluation, pre-extracted proposed "
        "tensors of shape (N,T,H,W,C)=(50,10,128,128,12) are used with spatial downsampling to "
        "32×32 for CPU-feasible re-training.",
    )

    H(doc, "4.3 SMA-CLMPNet Architecture", 2)
    body(
        doc,
        "Conv3D stacks with modified pooling reduce spatial extent while preserving strong "
        "activations. Distributed LSTM layers model temporal dependencies across frames. SCAM "
        "jointly reweights spatial locations and channels; MUSE applies parallel multiscale "
        "context. Fused SMA features pass dense layers and softmax classification. In code, the "
        "P1-Proposed implementation realizes ThreeDCNNLSTM with SCAM and MUSE (opt=3).",
    )

    H(doc, "4.4 Multi-Model Cohort for Fair Comparison", 2)
    body(
        doc,
        "For transparent comparison on identical index splits: DCNN (compact Conv2D), "
        "EfficientNetV2B0 (latest TF-2.10 backbone), MobileNetV2 (mobile CNN), STIDNet-style "
        "teacher–student distillation, and P1-Proposed (SMA-CLMPNet path).",
    )

    # 5 Results GENUINE ONLY
    H(doc, "5. Results (Genuine Machine Measurements Only)", 1)
    P(
        doc,
        "IMPORTANT: All numerical tables in this section are GENUINE. They were produced by "
        "Optimized/evaluate_multi.py using Optimized/metrics_fixed.py (sklearn). Fabricated "
        "Analysis/TP/COM_A.npy values (~93–96% Acc from the mealpy metric path) are NOT reported.",
        size=10, bold=True, after=8, color=RGBColor(0x8B, 0x00, 0x00),
    )

    H(doc, "5.1 Dataset and Protocol", 2)
    body(
        doc,
        "Local feature cache: Features/Features.pkl, proposed tensor (50, 10, 128, 128, 12), "
        "labels 29 authentic / 21 forged (majority baseline accuracy 58%). Train fractions 80% "
        "and 90% (seed-stratified). Test sizes ≈10 and ≈5 induce high variance; therefore "
        "Balanced Accuracy and F1 are preferred over raw Accuracy when class collapse is present "
        "(Sen≈0, Spec≈1).",
    )

    H(doc, "5.2 Performance Metrics", 2)
    body(
        doc,
        "Accuracy, sensitivity (recall of forged class), specificity, precision, F1-score, and "
        "balanced accuracy are computed from the confusion matrix with sklearn — never mealpy.metrics.",
    )

    models = ["DCNN", "EfficientNetV2B0", "MobileNetV2", "STIDNet", "P1-Proposed"]

    # Fresh genuine re-runs (primary)
    g1 = RES / "evaluation_multi_ep20_genuine.csv"
    g2 = RES / "evaluation_multi_ep20_genuine_bal_cw_os.csv"
    if g1.exists():
        metrics_table(
            doc, read_csv(g1), models,
            "5.3 Multi-model results — epochs=20 (fresh genuine re-run, no balance)",
        )
        cap(doc, "Table. GENUINE metrics (%) at 80% and 90% train; source evaluation_multi_ep20_genuine.csv")
    if g2.exists():
        metrics_table(
            doc, read_csv(g2), models,
            "5.4 Multi-model results — epochs=20 + class_weight + oversample (fresh genuine re-run)",
        )
        cap(doc, "Table. GENUINE balanced-training metrics; source evaluation_multi_ep20_genuine_bal_cw_os.csv")

    H(doc, "5.5 Additional Genuine Archive Ladder", 2)
    body(
        doc,
        "Earlier genuine sklearn re-runs on the same feature cache (same metric path) are retained "
        "for transparency of the optimization campaign.",
    )
    for label, name in [
        ("epochs=2", "evaluation_multi_ep2.csv"),
        ("epochs=20 (prior)", "evaluation_multi_ep20.csv"),
        ("epochs=50", "evaluation_multi_ep50.csv"),
        ("epochs=100", "evaluation_multi_ep100.csv"),
        ("epochs=20 + balance (prior)", "evaluation_multi_ep20_bal_cw_os.csv"),
    ]:
        p = RES / name
        if p.exists():
            metrics_table(doc, read_csv(p), models, f"Archive — {label}")

    H(doc, "5.6 Figures (genuine multi-model plots)", 2)
    add_img(doc, FIGS / "Fig1_machine_accuracy.png",
            caption="Figure. Multi-model accuracy at 80% and 90% training (genuine run).")
    add_img(doc, FIGS / "Fig_metrics_80.png",
            caption="Figure. Metric suite @80% train (genuine).")
    add_img(doc, FIGS / "Fig_metrics_90.png",
            caption="Figure. Metric suite @90% train (genuine).")
    add_img(doc, FIGS / "Fig_ranking_last_split.png",
            caption="Figure. Model ranking by accuracy on last split (genuine).")

    H(doc, "5.7 Comparative Discussion", 2)
    body(
        doc,
        "Under genuine metrics on N=50, no model reaches the 95–99% accuracy band sometimes "
        "quoted from corrupted evaluation paths. The best non-collapse accuracy in the fresh "
        "re-run is 80% (DCNN @90%, F1=0.67, BalAcc=0.75). With class balancing, EfficientNetV2B0 "
        "attains Acc 80% and BalAcc 68.8% @80% with F1=0.50, indicating improved minority "
        "sensitivity relative to pure majority predictors. P1-Proposed remains variance-limited "
        "on tiny test folds; longer full-protocol training on complete FaceForensics++ video sets "
        "is required for production-grade claims.",
    )

    # 6 Conclusion
    H(doc, "6. Conclusion", 1)
    body(
        doc,
        "This paper presented SMA-CLMPNet for intra-frame video forgery detection following the "
        "Paper 1 methodological template, and reported a complete machine-verified multi-model "
        "study using only genuine sklearn metrics. Fabricated Analysis/COM_A and related TP/KF "
        "plots have been removed from the working repository. Results support transparent "
        "comparison among modern backbones and the proposed attention–LSTM hybrid under a "
        "controlled feature cache, and motivate larger-scale data and k-fold protocols for "
        "future SCI dissemination under the RIITC–NKUST postdoctoral programme.",
    )

    H(doc, "7. Future Work", 1)
    body(
        doc,
        "(i) Full FaceForensics++ re-extraction and longer schedules; (ii) stratified k-fold CV; "
        "(iii) class-balanced losses / focal loss; (iv) multi-modal fusion toward document/news "
        "streams for semantic media forensics; (v) SCI journal submission with genuine-only tables.",
    )

    H(doc, "Acknowledgment", 1)
    body(
        doc,
        "The authors thank NKUST RIITC (Ref. RIITC-Postdoc-2027-B03), Chitkara University, and "
        "VIPS-TC for academic support.",
    )

    H(doc, "References (DOI / peer-reviewed sources used in discussion)", 1)
    refs = [
        '[1] A. Rössler et al., "FaceForensics++: Learning to detect manipulated facial images," in Proc. IEEE/CVF ICCV, 2019. DOI: 10.1109/ICCV.2019.00009.',
        '[2] M. Tan and Q. V. Le, "EfficientNetV2: Smaller models and faster training," in Proc. ICML, 2021. DOI: 10.48550/arXiv.2104.00298.',
        '[3] M. Sandler et al., "MobileNetV2: Inverted residuals and linear bottlenecks," in Proc. IEEE/CVF CVPR, 2018. DOI: 10.1109/CVPR.2018.00474.',
        '[4] I. J. Goodfellow et al., "Generative adversarial nets," in Proc. NeurIPS, 2014. DOI: 10.48550/arXiv.1406.2661.',
        '[5] H. Farid, "Image forgery detection," IEEE Signal Process. Mag., vol. 26, no. 2, pp. 16–25, 2009. DOI: 10.1109/MSP.2008.931079.',
        '[6] O. Ronneberger, P. Fischer, and T. Brox, "U-Net: Convolutional networks for biomedical image segmentation," in MICCAI, 2015. DOI: 10.1007/978-3-319-24574-4_28.',
        '[7] V. Christlein et al., "An evaluation of popular copy-move forgery detection approaches," IEEE Trans. Inf. Forensics Security, 2012. DOI: 10.1109/TIFS.2012.2218597.',
        '[8] L. D\'Amiano et al., "A PatchMatch-based dense-field algorithm for video copy–move detection and localization," IEEE Trans. Circuits Syst. Video Technol., 2019. DOI: 10.1109/TCSVT.2018.2804768.',
        '[9] S. Lyu, X. Pan, and X. Zhang, "Exposing region splicing forgeries with blind local noise estimation," Int. J. Comput. Vis., 2014. DOI: 10.1007/s11263-013-0688-y.',
        '[10] T. J. de Carvalho et al., "Exposing digital image forgeries by illumination color classification," IEEE Trans. Inf. Forensics Security, 2013. DOI: 10.1109/TIFS.2013.2265677.',
        '[11] Machine multi-model logs: Optimized/results/evaluation_multi_ep20_genuine*.csv (sklearn metrics_fixed; 2026-08-06 re-run).',
        '[12] Integrity note: Optimized/INTEGRITY_FINDING.md — mealpy.metrics path is not used for reported tables.',
    ]
    for r in refs:
        P(doc, r, size=9, align="justify", after=4)

    try:
        doc.save(str(OUT))
        saved = OUT
    except PermissionError:
        saved = OPT / "Paper1_SMA_CLMPNet_Genuine_Research_Paper_v2.docx"
        doc.save(str(saved))
    try:
        import shutil
        shutil.copy2(saved, OUT2)
    except Exception as e:
        print("copy root failed", e)
    # also replace older genuine article name for consistency
    try:
        import shutil
        shutil.copy2(saved, OPT / "Paper1_Genuine_Results_Article.docx")
        shutil.copy2(saved, ROOT / "Paper1_Genuine_Results_Article.docx")
    except Exception:
        pass
    print("SAVED", saved, saved.stat().st_size)
    return saved


if __name__ == "__main__":
    build()
