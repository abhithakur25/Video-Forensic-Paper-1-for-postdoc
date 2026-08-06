# -*- coding: utf-8 -*-
"""Build Paper 1 article with GENUINE results only (sklearn metrics_fixed)."""
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\USER\Downloads\PostDoc\CODE_28-04-2025_Paper1")
OPT = ROOT / "Optimized"
RES = OPT / "results"
FIGS = OPT / "figures"
OUT = OPT / "Paper1_Genuine_Results_Article.docx"
OUT_ROOT = ROOT / "Paper1_Genuine_Results_Article.docx"


def set_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def para(doc, text="", size=11, bold=False, italic=False, align="left",
         before=0, after=6, first=0, color=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first:
        pf.first_line_indent = Inches(first)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = para(doc, "", size=sizes.get(level, 11), before=12 if level == 1 else 8, after=4)
    r = p.add_run(text)
    set_run(r, size=sizes.get(level, 11), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(cell, "1F4E79")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run(r, size=9)
            if i % 2:
                shade(cell, "E8F0F8")
    para(doc, "", after=4)
    return t


def add_img(doc, path, width=5.5, cap=None):
    path = Path(path)
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    if cap:
        para(doc, cap, size=9, italic=True, align="center", after=10)


def read_csv(path):
    out = {}
    path = Path(path)
    if not path.exists():
        return out
    with open(path, newline="", encoding="utf-8") as fh:
        for row in csv.DictReader(fh):
            m = row["model"]
            tp = int(float(row["train_pct"]))
            out.setdefault(m, {})[tp] = row
    return out


def pct(v):
    try:
        x = float(v)
        return f"{x * 100:.2f}%"
    except Exception:
        return "—"


def flag_row(r):
    """Mark majority collapse on genuine rows."""
    try:
        sen, spe = float(r.get("SEN", 0)), float(r.get("SPE", 0))
        if (sen < 0.05 and spe > 0.95) or (spe < 0.05 and sen > 0.95):
            return " †"
    except Exception:
        pass
    return ""


def metrics_block(doc, data, models, title):
    heading(doc, title, 2)
    para(
        doc,
        "TAG: GENUINE — sklearn metrics_fixed.py (not mealpy). "
        "† = class collapse (Acc misleading; prefer BalAcc/F1).",
        size=9, italic=True, after=6,
    )
    rows = []
    for m in models:
        r80, r90 = data.get(m, {}).get(80, {}), data.get(m, {}).get(90, {})
        rows.append([
            m + flag_row(r80),
            pct(r80.get("ACC")), pct(r80.get("SEN")), pct(r80.get("SPE")),
            pct(r80.get("PRE")), pct(r80.get("F1")), pct(r80.get("BAL_ACC")),
            pct(r90.get("ACC")), pct(r90.get("SEN")), pct(r90.get("SPE")),
            pct(r90.get("PRE")), pct(r90.get("F1")), pct(r90.get("BAL_ACC")),
        ])
    table(
        doc,
        ["Model", "A80", "Se80", "Sp80", "P80", "F180", "B80",
         "A90", "Se90", "Sp90", "P90", "F190", "B90"],
        rows,
    )


def find_csv(patterns):
    """Return first existing results CSV matching any pattern substring."""
    files = sorted(RES.glob("evaluation_multi_*.csv"), key=lambda p: p.stat().st_mtime, reverse=True)
    for pat in patterns:
        for f in files:
            if pat in f.name:
                return f
    return None


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    para(doc, "Paper 1 — Genuine Multi-Model Evaluation Article", size=16, bold=True, align="center", after=4)
    para(
        doc,
        "Video forgery detection with attention (LDZP + optical flow) — "
        "CODE_28-04-2025_Paper1 / Optimized package",
        size=11, italic=True, align="center", after=4,
    )
    para(
        doc,
        "ALL TABLES BELOW ARE GENUINE. Fabricated mealpy-based Analysis/COM_A and "
        "Results/TP–KF artefacts have been removed from the working tree and are not cited.",
        size=10, bold=True, align="center", after=8, color=RGBColor(0x8B, 0x00, 0x00),
    )
    para(
        doc,
        "Authors (Offer RIITC-Postdoc-2027-B03): Dr. Abhishek Thakur (Chitkara University / NKUST RIITC); "
        "Prof. Vishal Jain (VIPS-TC, co-supervisor); Prof. Chin-Shiuh Shieh (NKUST RIITC, supervisor). "
        "Report date: 2026-08-06.",
        size=9, align="center", after=12,
    )

    heading(doc, "1. Integrity statement", 1)
    para(
        doc,
        "Original SubFunctions/Evaluate.py scores through mealpy.metrics.confusion_matrix, which "
        "discards model predictions and injects random label flips. That path is untrustworthy. "
        "This article reports only scores from Optimized/metrics_fixed.py (sklearn). "
        "Fabricated NPY folders (Analysis/, Analysis1/) and plots derived from them "
        "(Results/TP, Results/KF) are excluded from this package.",
        align="justify", first=0.25,
    )

    heading(doc, "2. Experimental setup (genuine)", 1)
    table(
        doc,
        ["Item", "Value"],
        [
            ["Features", "Features/Features.pkl → proposed (50, 10, 128, 128, 12)"],
            ["Labels", "29 authentic / 21 forged (majority Acc baseline 58%)"],
            ["Train fractions", "80% and 90% (seed-stratified)"],
            ["Test size", "~10 @80%, ~5 @90% (high variance)"],
            ["Models", "DCNN, EfficientNetV2B0, MobileNetV2, STIDNet, P1-Proposed"],
            ["Proposed architecture", "3D-CNN + dual LSTM + SCAM + MUSE (opt=3)"],
            ["Metrics", "Acc, Sen, Spec, Prec, F1, BalAcc (sklearn)"],
            ["Env", "VideoForgeryCPU, TensorFlow/Keras 2.10"],
        ],
    )

    models = ["DCNN", "EfficientNetV2B0", "MobileNetV2", "STIDNet", "P1-Proposed"]

    # Prefer freshly tagged genuine runs; fall back to prior ep20 / bal files
    csv_base = find_csv(["ep20_genuine", "ep20.csv"]) or RES / "evaluation_multi_ep20.csv"
    # note: ep20.csv alone might match ep20_bal - order by specificity
    candidates_base = list(RES.glob("evaluation_multi_ep20_genuine*.csv"))
    candidates_base = [c for c in candidates_base if "bal" not in c.name]
    if not candidates_base:
        candidates_base = [RES / "evaluation_multi_ep20.csv"] if (RES / "evaluation_multi_ep20.csv").exists() else []
    candidates_bal = list(RES.glob("evaluation_multi_ep20*bal*.csv")) + list(
        RES.glob("evaluation_multi_ep20_genuine_bal*.csv")
    )
    if not candidates_bal:
        candidates_bal = list(RES.glob("evaluation_multi_ep20_bal*.csv"))

    csv_base = candidates_base[0] if candidates_base else None
    csv_bal = candidates_bal[0] if candidates_bal else None

    heading(doc, "3. Genuine results — epochs = 20 (no balance)", 1)
    if csv_base:
        para(doc, f"Source file: {csv_base.name}", size=9, italic=True)
        metrics_block(doc, read_csv(csv_base), models, "Table — full metrics @80% and @90%")
    else:
        para(doc, "No baseline genuine CSV found.", italic=True)

    heading(doc, "4. Genuine results — epochs = 20 + class_weight + oversample", 1)
    if csv_bal:
        para(doc, f"Source file: {csv_bal.name}", size=9, italic=True)
        metrics_block(doc, read_csv(csv_bal), models, "Table — balanced training metrics")
    else:
        para(doc, "No balanced genuine CSV found.", italic=True)

    # Prior genuine archive runs (if still present after cleanup of fabrications only)
    heading(doc, "5. Additional genuine archive runs (epochs ladder)", 1)
    para(
        doc,
        "Earlier genuine re-runs with the same sklearn path (before this final clean-up). "
        "Included only if CSV files remain under Optimized/results/.",
        size=10, after=6,
    )
    for label, fname in [
        ("epochs=2", "evaluation_multi_ep2.csv"),
        ("epochs=50", "evaluation_multi_ep50.csv"),
        ("epochs=100", "evaluation_multi_ep100.csv"),
    ]:
        p = RES / fname
        if p.exists():
            metrics_block(doc, read_csv(p), models, f"Table — {label}")

    heading(doc, "6. Best genuine scores (summary)", 1)
    # compute from available
    best = {"acc": (0, ""), "bal": (0, ""), "f1": (0, "")}
    for f in RES.glob("evaluation_multi_*.csv"):
        if "backup" in str(f):
            continue
        data = read_csv(f)
        for m, tps in data.items():
            for tp, r in tps.items():
                try:
                    a, b, f1 = float(r["ACC"]), float(r["BAL_ACC"]), float(r["F1"])
                except Exception:
                    continue
                tag = f"{m} | {f.stem} | {tp}%"
                # prefer non-collapse for "best acc"
                sen, spe = float(r.get("SEN", 0)), float(r.get("SPE", 0))
                collapse = (sen < 0.05 and spe > 0.95) or (spe < 0.05 and sen > 0.95)
                if not collapse and a > best["acc"][0]:
                    best["acc"] = (a, tag)
                if b > best["bal"][0]:
                    best["bal"] = (b, tag)
                if f1 > best["f1"][0]:
                    best["f1"] = (f1, tag)
    table(
        doc,
        ["Metric", "Best genuine value", "Where"],
        [
            ["Accuracy (non-collapse)", f"{best['acc'][0]*100:.1f}%", best["acc"][1]],
            ["Balanced accuracy", f"{best['bal'][0]*100:.1f}%", best["bal"][1]],
            ["F1-score", f"{best['f1'][0]*100:.1f}%", best["f1"][1]],
            ["95–99% Acc target", "NOT achieved", "N=50 + honest metrics"],
        ],
    )

    heading(doc, "7. Figures (genuine multi-model plots)", 1)
    add_img(doc, FIGS / "Fig1_machine_accuracy.png", cap="Fig. 1. Multi-model accuracy (genuine run).")
    add_img(doc, FIGS / "Fig_metrics_80.png", cap="Fig. 2. Metrics @80% train (genuine).")
    add_img(doc, FIGS / "Fig_metrics_90.png", cap="Fig. 3. Metrics @90% train (genuine).")
    add_img(doc, FIGS / "Fig_ranking_last_split.png", cap="Fig. 4. Ranking by accuracy (genuine).")

    heading(doc, "8. Reproducibility", 1)
    for line in [
        r'$E = "C:\Users\USER\anaconda3\envs\VideoForgeryCPU"',
        r'$env:PATH = "$E\Library\bin;$E;$E\Scripts;" + $env:PATH',
        r'cd CODE_28-04-2025_Paper1',
        r"# Place Features.pkl under Features/",
        r'& "$E\python.exe" -u Optimized\evaluate_multi.py --epochs 20 --train-pcts "0.8,0.9" --tag genuine',
        r'& "$E\python.exe" -u Optimized\evaluate_multi.py --epochs 20 --class-weight --oversample --tag genuine_bal',
    ]:
        para(doc, line, size=8, after=2)

    heading(doc, "9. Conclusion", 1)
    para(
        doc,
        "This article retains only genuine, machine-measured multi-model results under honest "
        "sklearn metrics. Fabricated COM_A / Analysis / TP–KF paper artefacts have been removed. "
        "On the 50-sample feature cache, best genuine BalAcc is about 81% (STIDNet or MobileNetV2 "
        "with balance); 95–99% test accuracy is not supported by genuine measurements.",
        align="justify", first=0.25,
    )

    try:
        doc.save(str(OUT))
        saved = OUT
    except PermissionError:
        saved = OPT / "Paper1_Genuine_Results_Article_v2.docx"
        doc.save(str(saved))
    try:
        import shutil
        shutil.copy2(saved, OUT_ROOT)
    except Exception:
        pass
    print("SAVED", saved)
    return saved


if __name__ == "__main__":
    build()
