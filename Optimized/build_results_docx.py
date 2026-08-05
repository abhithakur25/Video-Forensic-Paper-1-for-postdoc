# -*- coding: utf-8 -*-
"""Build comprehensive Paper 1 evaluation report DOCX with every step + all results."""
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\USER\Downloads\PostDoc\CODE_28-04-2025_Paper1")
OPT = ROOT / "Optimized"
OUT = OPT / "Paper1_MultiModel_Evaluation_Full_Report.docx"
OUT_ALT = OPT / "Paper1_MultiModel_Evaluation_Full_Report_v2.docx"
FIGS = OPT / "figures"
RES = OPT / "results"


def set_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def para(doc, text="", size=11, bold=False, italic=False, align="left",
         before=0, after=6, first=0, color=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first:
        pf.first_line_indent = Inches(first)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = para(doc, "", size=sizes.get(level, 11), before=12 if level == 1 else 8, after=4)
    r = p.add_run(text)
    set_run(r, size=sizes.get(level, 11), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(cell, "1F4E79")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run(r, size=9)
            if i % 2:
                shade(cell, "E8F0F8")
    para(doc, "", after=4)
    return t


def add_img(doc, path, width=5.5, cap=None):
    path = Path(path)
    if not path.exists():
        para(doc, f"[Figure missing: {path.name}]", size=9, italic=True, align="center")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    if cap:
        para(doc, cap, size=9, italic=True, align="center", after=10)


def read_csv_metrics(csv_path):
    """Return dict model -> {80: {...}, 90: {...}}"""
    out = {}
    if not Path(csv_path).exists():
        return out
    with open(csv_path, newline="", encoding="utf-8") as fh:
        for row in csv.DictReader(fh):
            m = row["model"]
            tp = int(float(row["train_pct"]))
            out.setdefault(m, {})[tp] = row
    return out


def pct(v):
    try:
        return f"{float(v)*100:.2f}%"
    except Exception:
        return str(v) if v else "—"


def metrics_table(doc, data, models, title):
    heading(doc, title, 2)
    # Accuracy
    para(doc, "Accuracy", size=10, bold=True, after=2)
    rows = []
    for m in models:
        r80 = data.get(m, {}).get(80, {})
        r90 = data.get(m, {}).get(90, {})
        rows.append([m, pct(r80.get("ACC", "")), pct(r90.get("ACC", ""))])
    table(doc, ["Model", "Acc @80%", "Acc @90%"], rows)

    para(doc, "Full metrics @80% training", size=10, bold=True, after=2)
    rows = []
    for m in models:
        r = data.get(m, {}).get(80, {})
        rows.append([
            m,
            pct(r.get("ACC", "")),
            pct(r.get("SEN", "")),
            pct(r.get("SPE", "")),
            pct(r.get("PRE", "")),
            pct(r.get("F1", "")),
            pct(r.get("BAL_ACC", r.get("BAL", ""))),
        ])
    table(doc, ["Model", "Acc", "Sen", "Spec", "Prec", "F1", "BalAcc"], rows)

    para(doc, "Full metrics @90% training", size=10, bold=True, after=2)
    rows = []
    for m in models:
        r = data.get(m, {}).get(90, {})
        rows.append([
            m,
            pct(r.get("ACC", "")),
            pct(r.get("SEN", "")),
            pct(r.get("SPE", "")),
            pct(r.get("PRE", "")),
            pct(r.get("F1", "")),
            pct(r.get("BAL_ACC", r.get("BAL", ""))),
        ])
    table(doc, ["Model", "Acc", "Sen", "Spec", "Prec", "F1", "BalAcc"], rows)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    para(
        doc,
        "Paper 1 — Multi-Model Optimization & Evaluation Report",
        size=16, bold=True, align="center", after=4,
    )
    para(
        doc,
        "Intra-frame / attention-based video forgery detection "
        "(LDZP + optical flow codebase: CODE_28-04-2025_Paper1)",
        size=11, italic=True, align="center", after=4,
    )
    para(
        doc,
        "Aligned with Paper 2 multi-model protocol (DCNN, EfficientNetV2B0, MobileNetV2, "
        "proposed model). All machine metrics use honest sklearn scoring "
        "(Optimized/metrics_fixed.py). Original mealpy.metrics path is documented as tampered.",
        size=10, align="center", after=8,
    )
    para(
        doc,
        "Report date: 2026-08-05  |  Environment: VideoForgeryCPU (TF/Keras 2.10)  |  "
        "Output folder: CODE_28-04-2025_Paper1/Optimized/",
        size=9, align="center", after=12, color=RGBColor(0x44, 0x44, 0x44),
    )

    # ===== 1. Overview =====
    heading(doc, "1. Overview and Objectives", 1)
    para(
        doc,
        "After completing Paper 2 (OM²AHL-BiG multi-model evaluation), Paper 1 was optimized "
        "with the same methodology: trainable multi-model comparison on Features.pkl, honest "
        "metrics, result artefacts (CSV/TXT/NPY/figures), and a self-contained Optimized/ package. "
        "This document records every step executed and every result table generated.",
        align="justify", first=0.25,
    )
    para(doc, "Goals completed:", bold=True, after=2)
    for b in [
        "Explore Paper 1 codebase and Claude skill findings (metric integrity issue).",
        "Implement Optimized/ package: metrics_fixed, feature_adapters, MultiModel, evaluate_multi, balance.",
        "Run multi-model evaluation for epochs = 2, 20, 50, 100.",
        "Run class_weight + oversampling at epochs = 20.",
        "Tune minority class-weight scale (partial/completed sweep log).",
        "Save working code, optimized code, logs, figures, and this consolidated report.",
    ]:
        para(doc, "• " + b, size=10, after=2)

    # ===== 2. Integrity =====
    heading(doc, "2. Integrity Finding (Critical)", 1)
    para(
        doc,
        "Claude’s Paper 1 skill documented that vendored mealpy/metrics.py discards model "
        "predictions inside confusion_matrix and injects ground-truth labels plus random flips "
        "(per ≈ uniform 0.09–0.45). Therefore original Analysis/*.npy and figures regenerated "
        "via SubFunctions/Evaluate.py are not trustworthy model scores.",
        align="justify", first=0.25,
    )
    para(
        doc,
        "Fix: Optimized/metrics_fixed.py uses sklearn only (accuracy, balanced accuracy, "
        "precision, recall, F1, specificity). All tables below are from this path.",
        align="justify", first=0.25,
    )
    para(
        doc,
        "Reference only: Analysis/TP/COM_A.npy last-row accuracy ≈ 93.34% (paper artefact; "
        "may be affected by tampered metrics).",
        size=10, italic=True, after=8,
    )

    # ===== 3. Dataset =====
    heading(doc, "3. Dataset and Feature Cache", 1)
    table(
        doc,
        ["Item", "Value"],
        [
            ["File", "Features/Features.pkl"],
            ["Proposed tensor", "(50, 10, 128, 128, 12) — T×H×W×C"],
            ["Comparative tensors", "comparative1–5 also present (not primary in multi-model)"],
            ["Labels", "50 samples — class 0: 29, class 1: 21"],
            ["Majority baseline Acc", "58.00%"],
            ["Train splits used", "80% and 90% (seed-stratified)"],
            ["Test size @80% / @90%", "n=10 / n=5 (high variance)"],
            ["Spatial downsample", "32×32 for CPU feasibility (3D-CNN path)"],
        ],
    )

    # ===== 4. Models =====
    heading(doc, "4. Models Evaluated", 1)
    table(
        doc,
        ["Model", "Description", "Role"],
        [
            ["DCNN", "Compact Conv2D on time-mean spatial maps", "Baseline CNN"],
            ["EfficientNetV2B0", "Latest TF-2.10 keras.applications backbone", "Modern SOTA-style"],
            ["MobileNetV2", "Inverted residual mobile CNN", "Efficiency baseline"],
            ["STIDNet", "Teacher–student distillation (Paper 1 comparative)", "Paper comparative"],
            ["P1-Proposed", "3D-CNN + dual LSTM + SCAM + MUSE (ThreeDCNNLSTM opt=3)", "Proposed"],
            ["RF / GBM", "Classical ensembles on flattened features (weight-tune phase)", "Small-N boosters"],
        ],
    )

    # ===== 5. Package layout =====
    heading(doc, "5. Optimized Package Layout", 1)
    para(doc, "All artefacts live under CODE_28-04-2025_Paper1/Optimized/:", size=10)
    for line in [
        "evaluate_multi.py — multi-model driver (Paper-2 style)",
        "MultiModel.py — trainable models + RF/GBM",
        "metrics_fixed.py — honest sklearn metrics",
        "feature_adapters.py — tensor reshape/resize",
        "balance.py — class_weight + oversample + minority_scale",
        "run_weight_tuned.py — minority weight scale sweep",
        "INTEGRITY_FINDING.md, RESULTS_SUMMARY.md",
        "results/ — CSV, TXT, NPY for every run",
        "figures/ — accuracy / metrics / ranking charts (no on-image titles)",
        "working_code/ — original Model/MUSE/SCAM + optimized copies + tampered mealpy metrics archive",
        "logs/ — execution logs",
    ]:
        para(doc, "• " + line, size=9, after=1)

    # ===== 6. Step-by-step protocol =====
    heading(doc, "6. Step-by-Step Execution Protocol", 1)

    heading(doc, "Step 1 — Environment and integrity setup", 2)
    para(
        doc,
        "Conda env VideoForgeryCPU (Python 3.8, TensorFlow/Keras 2.10). PATH includes "
        "Library\\bin for scipy/skimage. Metrics forced to Optimized/metrics_fixed.py. "
        "Integrity note archived in INTEGRITY_FINDING.md and working_code/mealpy_metrics_TAMPERED.py.",
        align="justify", first=0.2,
    )

    heading(doc, "Step 2 — Baseline multi-model run (epochs = 2)", 2)
    para(
        doc,
        "Command: python -u Optimized/evaluate_multi.py --epochs 2 --train-pcts 0.8,0.9. "
        "Wall ≈ 93 s. Outputs: evaluation_multi_ep2.txt/csv, MULTI_*.npy, figures.",
        align="justify", first=0.2,
    )
    data_ep2 = read_csv_metrics(RES / "evaluation_multi_ep2.csv")
    models5 = ["DCNN", "EfficientNetV2B0", "MobileNetV2", "STIDNet", "P1-Proposed"]
    metrics_table(doc, data_ep2, models5, "Step 2 results — epochs = 2")

    heading(doc, "Step 3 — Increased epochs (epochs = 20)", 2)
    para(
        doc,
        "User choice: re-run with more epochs. Command: --epochs 20. Wall ≈ 232 s. "
        "Outputs: evaluation_multi_ep20.*",
        align="justify", first=0.2,
    )
    data_ep20 = read_csv_metrics(RES / "evaluation_multi_ep20.csv")
    metrics_table(doc, data_ep20, models5, "Step 3 results — epochs = 20")

    heading(doc, "Step 4 — Further epochs (epochs = 50)", 2)
    para(
        doc,
        "User choice: epochs = 50. Wall ≈ 484 s. Outputs: evaluation_multi_ep50.*",
        align="justify", first=0.2,
    )
    data_ep50 = read_csv_metrics(RES / "evaluation_multi_ep50.csv")
    metrics_table(doc, data_ep50, models5, "Step 4 results — epochs = 50")

    heading(doc, "Step 5 — High epoch budget (epochs = 100)", 2)
    para(
        doc,
        "User choice: epochs = 100. Wall ≈ 1060 s. Outputs: evaluation_multi_ep100.* "
        "Conclusion recorded: more epochs alone does not systematically improve BalAcc/F1; "
        "several high Acc scores are majority-class collapses (Sen=0, Spec=1).",
        align="justify", first=0.2,
    )
    data_ep100 = read_csv_metrics(RES / "evaluation_multi_ep100.csv")
    metrics_table(doc, data_ep100, models5, "Step 5 results — epochs = 100")

    heading(doc, "Step 6 — Class weights + oversampling (epochs = 20)", 2)
    para(
        doc,
        "User choice: balance mode. Command: --epochs 20 --class-weight --oversample --tag bal. "
        "Uses balanced class_weight in Keras fit + random minority oversample on train only. "
        "Wall ≈ 558 s. Outputs: evaluation_multi_ep20_bal_cw_os.*",
        align="justify", first=0.2,
    )
    data_bal = read_csv_metrics(RES / "evaluation_multi_ep20_bal_cw_os.csv")
    metrics_table(doc, data_bal, models5, "Step 6 results — epochs = 20 + class_weight + oversample")

    heading(doc, "Step 7 — Minority weight-scale tuning", 2)
    para(
        doc,
        "User requested weight adjustment aiming toward 95–99% accuracy. Implemented "
        "Optimized/balance.py minority_scale and Optimized/run_weight_tuned.py sweeping "
        "minority_scale ∈ {1.0, 2.0, 3.0, 4.0} and oversample_ratio ∈ {1.0, 1.5} for models "
        "DCNN, EfficientNetV2B0, MobileNetV2, STIDNet, P1-Proposed, RF, GBM (epochs_deep=30). "
        "Metrics remain sklearn-only (not mealpy). Best configurations from the sweep log:",
        align="justify", first=0.2,
    )
    table(
        doc,
        ["Model", "Best minority_scale", "Best oversample_ratio", "Mean BalAcc @80%", "Mean Acc @80%"],
        [
            ["DCNN", "3.0", "1.5", "68.8%", "50.0%"],
            ["EfficientNetV2B0", "1.0", "1.5", "81.2%", "70.0%"],
            ["MobileNetV2", "2.0", "1.0", "81.2%", "70.0%"],
            ["STIDNet", "3.0", "1.5", "68.8%", "80.0%"],
            ["P1-Proposed", "2.0", "1.0", "50.0%", "80.0%†"],
            ["RF", "1.0", "1.0", "62.5%", "70.0%"],
            ["GBM", "1.0 / 2.0", "1.5", "43.8%", "70.0%"],
        ],
    )
    para(
        doc,
        "† High Acc with BalAcc=50% typically indicates majority collapse. "
        "Target 95–99% test accuracy was NOT reached with honest metrics on N=50; "
        "paper ~93% COM_A remains reference-only.",
        size=10, italic=True, after=8,
    )

    # Weight sweep detail table from log
    heading(doc, "Step 7 detail — weight sweep log (selected rows)", 2)
    para(
        doc,
        "Full sweep is archived in Optimized/logs/weight_tune.log. Representative best-per-model "
        "rows are above. Sweep files (when completed): evaluation_weight_tuned.txt/csv, weight_sweep.csv.",
        size=10, after=6,
    )

    # ===== 7. Cross-run comparison =====
    heading(doc, "7. Cross-Run Comparison Summary", 1)
    para(doc, "Best accuracy / balanced accuracy / F1 observed per run (machine, honest metrics):", size=10)
    table(
        doc,
        ["Run", "Best Acc", "Best BalAcc", "Best F1", "Wall time"],
        [
            ["epochs=2", "80% (DCNN/P1 @80%; EffNet @90%)", "75% EffNet @90%", "0.67 EffNet @90%", "~93 s"],
            ["epochs=20", "80% DCNN @90%; 70% STID @80%", "81% STID @80%; 75% DCNN @90%", "0.67 DCNN @90%", "~232 s"],
            ["epochs=50", "80% EffNet @80%†", "~50% typical", "0.57 P1 @90%", "~484 s"],
            ["epochs=100", "80% EffNet/P1 @80%†", "75% MobileNet @80%", "0.50 MobileNet @80%", "~1060 s"],
            ["ep20 + cw + os", "80% STID @80%†; 70% MobileNet", "81% MobileNet @80%", "0.57 MobileNet @80%", "~558 s"],
            ["weight-scale tune", "80% several†; 70% MobileNet/EffNet", "81% MobileNet/EffNet", "— (sweep mean)", "long"],
        ],
    )
    para(
        doc,
        "† Majority-class collapse (Sen≈0, Spec≈1). Prefer BalAcc and F1 over Acc on this corpus.",
        size=9, italic=True, after=8,
    )

    # ===== 8. Figures =====
    heading(doc, "8. Figures (latest regenerate from last evaluate_multi run)", 1)
    add_img(doc, FIGS / "Fig1_machine_accuracy.png",
            cap="Figure 1. Multi-model accuracy at 80% and 90% training (latest run).")
    add_img(doc, FIGS / "Fig_metrics_80.png",
            cap="Figure 2. Full metric suite @80% training (latest run).")
    add_img(doc, FIGS / "Fig_metrics_90.png",
            cap="Figure 3. Full metric suite @90% training (latest run).")
    add_img(doc, FIGS / "Fig_ranking_last_split.png",
            cap="Figure 4. Model ranking by accuracy on last split (latest run).")

    # ===== 9. How to reproduce =====
    heading(doc, "9. How to Reproduce", 1)
    para(doc, "From project root CODE_28-04-2025_Paper1:", size=10)
    for cmd in [
        r'$E = "C:\Users\USER\anaconda3\envs\VideoForgeryCPU"',
        r'$env:PATH = "$E\Library\bin;$E;$E\Scripts;" + $env:PATH',
        r'cd C:\Users\USER\Downloads\PostDoc\CODE_28-04-2025_Paper1',
        r'& "$E\python.exe" -u Optimized\evaluate_multi.py --epochs 20 --train-pcts "0.8,0.9"',
        r'& "$E\python.exe" -u Optimized\evaluate_multi.py --epochs 20 --class-weight --oversample --tag bal',
        r'& "$E\python.exe" -u Optimized\run_weight_tuned.py',
    ]:
        para(doc, cmd, size=8, after=2)

    # ===== 10. Conclusions =====
    heading(doc, "10. Conclusions", 1)
    for c in [
        "Paper 1 multi-model evaluation package is complete under Optimized/, parallel to Paper 2.",
        "Honest metrics show majority baseline 58%; test n=5–10 causes high variance.",
        "Increasing epochs (2→100) does not reliably reach 95–99% test accuracy.",
        "Class weights + oversampling improve minority recall for some models (e.g. MobileNetV2 BalAcc 81%).",
        "Weight-scale tuning further identified best minority_scale per model; peak honest Acc/BalAcc remain ~70–81%, not 95–99%.",
        "Paper COM_A ≈ 93% is reference-only and may stem from the tampered mealpy metric path.",
        "Reaching 95–99% with honest metrics likely requires larger data (full FaceForensics++), k-fold CV, and longer full-protocol training — not weight adjustment alone on N=50.",
    ]:
        para(doc, "• " + c, size=10, after=3)

    heading(doc, "11. File Index of Saved Results", 1)
    files = sorted(RES.glob("*")) if RES.exists() else []
    rows = [[f.name, f.stat().st_size, f.suffix] for f in files]
    if rows:
        table(doc, ["Filename", "Bytes", "Type"], rows)
    else:
        para(doc, "No files in results/", size=10)

    para(
        doc,
        "End of report. Source package: CODE_28-04-2025_Paper1/Optimized/",
        size=9, italic=True, align="center", before=12,
    )

    try:
        doc.save(str(OUT))
        saved = OUT
    except PermissionError:
        doc.save(str(OUT_ALT))
        saved = OUT_ALT
    print("SAVED", saved, saved.stat().st_size)
    return saved


if __name__ == "__main__":
    build()
